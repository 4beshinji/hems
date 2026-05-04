"""
DOCX loader — .docx files via python-docx paragraph extraction.
"""

from pathlib import Path

from loguru import logger

from .base import BaseLoader, DocumentEntry

try:
    from docx import Document

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logger.warning("python-docx not installed — DOCX loading disabled")


class DocxLoader(BaseLoader):
    extensions = [".docx"]
    doc_type = "docx"

    def load(self, file_path: Path, rel_path: str, source_name: str) -> DocumentEntry | None:
        if not _HAS_DOCX:
            return None

        stat = file_path.stat()
        title = Path(rel_path).stem
        tags = ["docx"]
        metadata: dict = {}

        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            body = "\n".join(paragraphs)

            # Extract title from first heading or paragraph
            for p in doc.paragraphs:
                if p.style.name.startswith("Heading") and p.text.strip():
                    title = p.text.strip()
                    break

            metadata["paragraphs"] = len(paragraphs)

            # Extract core properties
            props = doc.core_properties
            if props.author:
                metadata["author"] = props.author
            if props.title:
                title = props.title
                metadata["title"] = props.title
        except Exception as e:
            logger.warning(f"DOCX extraction failed for {rel_path}: {e}")
            return None

        return DocumentEntry(
            path=rel_path,
            source_name=source_name,
            title=title,
            doc_type=self.doc_type,
            tags=tags,
            metadata=metadata,
            word_count=len(body.split()),
            modified_at=stat.st_mtime,
            content=body,
            body=body,
        )
